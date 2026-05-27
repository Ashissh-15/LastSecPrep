from docx import Document
import os

from loaders.image_extractor import (
    extract_image_description
)


def load_docx(file_path):

    documents = []

    doc = Document(file_path)


    # ==========================
    # Paragraph extraction
    # ==========================

    text = []

    for para in doc.paragraphs:

        para_text = para.text.strip()

        if para_text:

            text.append(
                para_text
            )


    if text:

        documents.append({

            "text":"\n".join(
                text
            ),

            "source":file_path,

            "page":1
        })


    # ==========================
    # Table extraction
    # ==========================

    for table_num, table in enumerate(
        doc.tables
    ):

        header = None

        semantic_rows = []


        for row_num, row in enumerate(
            table.rows
        ):

            row_data = []


            for cell in row.cells:

                row_data.append(
                    cell.text.strip()
                )


            # First row becomes headers
            if row_num == 0:

                header = row_data

                continue


            if header:

                row_text = []


                for i in range(

                    min(
                        len(header),
                        len(row_data)
                    )

                ):

                    row_text.append(

f"{header[i]} : {row_data[i]}"

                    )


                semantic_rows.append(

                    " | ".join(
                        row_text
                    )

                )


        if semantic_rows:

            documents.append({

                "text":
f"""
TABLE {table_num+1}

{chr(10).join(semantic_rows)}
""",

                "source":file_path,

                "page":"table"
            })


    # ==========================
    # Image extraction
    # ==========================

    rels = doc.part.rels

    image_count = 0


    for rel in rels:

        rel = rels[rel]


        if "image" in rel.target_ref:

            image_count += 1

            image = rel.target_part.blob


            image_name = (
                f"docx_img_{image_count}.png"
            )


            with open(
                image_name,
                "wb"
            ) as f:

                f.write(
                    image
                )


            description = (
                extract_image_description(
                    image_name
                )
            )


            if description.strip():

                documents.append({

                    "text":
f"""
IMAGE {image_count}

{description}
""",

                    "source":file_path,

                    "page":"image"
                })


            if os.path.exists(
                image_name
            ):

                os.remove(
                    image_name
                )


    return documents